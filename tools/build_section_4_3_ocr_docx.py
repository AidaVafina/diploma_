from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
ARTIFACTS_DIR = ROOT / "artifacts" / "section_4_3"
OUT_PATH = ROOT / "deliverables" / "section_4_3_ocr_expanded.docx"


def set_run_font(run, *, size: int = 14, bold: bool = False, italic: bool = False, name: str = "Times New Roman") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "5")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "AAB4C4")
        borders.append(tag)
    tbl_pr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True)


def add_paragraph(doc: Document, text: str, *, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_run_font(run)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=12, italic=True)


def add_picture(doc: Document, path: Path, *, width_cm: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def set_cell_text(cell, text: str, *, size: float = 10.5, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths_cm: list[float], *, font_size: float = 10.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(caption)
    set_run_font(run, size=12, italic=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E7EEF8")
        set_cell_text(cell, header, size=10, bold=True)
        cell.width = Cm(widths_cm[index])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=font_size)
            cells[index].width = Cm(widths_cm[index])

    doc.add_paragraph()


def add_algorithm(doc: Document, caption: str, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(caption)
    set_run_font(run, size=12, italic=True)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    table.columns[0].width = Cm(15.8)
    set_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FAFBFC")
    cell.width = Cm(15.8)
    cell.text = ""

    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        set_run_font(run, size=10, name="Courier New")

    doc.add_paragraph()


def main() -> None:
    manifest = json.loads((ARTIFACTS_DIR / "manifest.json").read_text(encoding="utf-8"))
    page_one = manifest["pages"]["tome1_p01_orig"]
    page_two = manifest["pages"]["tome1_p02_orig"]
    selected_rows = manifest["selected_table_rows"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)

    add_heading(doc, "4.3. Модуль OCR текстовых блоков")

    add_paragraph(
        doc,
        "Модуль OCR текстовых блоков предназначен для распознавания только тех областей страницы, которые на этапе layout-анализа были отнесены к текстовому маршруту обработки. В текущей реализации основная логика сосредоточена в сервисе text_block_processor.py и вызывается из общего PDF-конвейера после получения списка routed-блоков. Такое решение позволяет не выполнять повторное распознавание для формул, таблиц и иллюстраций, а сосредоточить вычислительные ресурсы на тех фрагментах, из которых затем формируется читаемый текст страницы.",
    )
    add_paragraph(
        doc,
        "На вход модуль получает RGB-изображение страницы и список объектов Block. Каждый блок содержит идентификатор block_id, тип области, координаты bounding box, порядок чтения reading_order, маршрут route_to, а также предварительные подсказки seed_text, seed_latex и seed_confidence, сформированные на предыдущем этапе. На выходе строится объект PageContent, включающий page_text, presentation-представление страницы, список ProcessedBlock, количество блоков needs_review и путь к сериализованному JSON-результату.",
    )

    add_table(
        doc,
        "Таблица 5 — Ключевые структуры данных модуля OCR текстовых блоков",
        ["Сущность", "Ключевые поля", "Роль в модуле", "Формирование"],
        [
            [
                "Block",
                "block_id, type, bbox, reading_order, route_to, seed_text",
                "Описывает routed-область страницы, подготовленную после layout-анализа.",
                "build_routed_blocks_from_layout",
            ],
            [
                "OCRResult",
                "text, confidence",
                "Хранит результат распознавания одного crop-фрагмента.",
                "recognize_text_block или fallback_text_result",
            ],
            [
                "ProcessedBlock",
                "content, confidence, needs_review, crop_data_url, ocr_backend",
                "Фиксирует итог обработки конкретного блока и метаданные контроля качества.",
                "process_text_blocks",
            ],
            [
                "PageContent",
                "page_content_id, page_text, presentation, blocks, result_json_path",
                "Объединяет все блоки страницы в общий структурированный результат.",
                "build_page_content + save_results",
            ],
        ],
        [2.4, 4.0, 5.4, 4.0],
        font_size=10.0,
    )

    add_paragraph(
        doc,
        "После загрузки изображения вызывается функция sort_blocks, которая упорядочивает routed-блоки по четырем признакам: reading_order, верхней координате y, левой координате x и идентификатору блока. Такой механизм устраняет неоднозначность при близких значениях координат и обеспечивает устойчивое восстановление порядка чтения даже в случае сложной верстки. Затем для каждого блока дополнительно применяется clamp_bbox, ограничивающий координаты границами изображения, и только после этого выполняется вырезание crop-фрагмента.",
    )

    add_picture(
        doc,
        ARTIFACTS_DIR / "ocr_pipeline_scheme.png",
        width_cm=15.6,
        caption="Рисунок 6 — Последовательность обработки routed-блоков в модуле OCR текстовых блоков",
    )

    add_paragraph(
        doc,
        "Распознавание текста реализовано через функцию recognize_text_block. Модуль пытается использовать PaddleOCR по двум совместимым сценариям API: сначала вызывается метод predict, характерный для более новых версий библиотеки, а при отсутствии нужного интерфейса выполняется переход к классическому методу ocr. Из ответа извлекаются распознанные строки и confidence, после чего формируется единый объект OCRResult. Благодаря такому решению сервис не привязан к одной конкретной версии PaddleOCR и допускает более гибкое развёртывание.",
    )
    add_paragraph(
        doc,
        "Если backend PaddleOCR недоступен или завершился с ошибкой, выполняется резервный сценарий fallback_text_result: в качестве текста используется seed_text, уже полученный на этапе Surya layout/OCR, а в качестве confidence — сохранённое seed_confidence. Параллельно поле ocr_backend получает значение surya_seed. Данный механизм предотвращает потерю содержимого страницы и позволяет конвейеру завершить обработку документа даже при частичной недоступности OCR-инфраструктуры.",
    )
    add_paragraph(
        doc,
        "Отдельное место в реализации занимает постобработка текста. Функция postprocess_text удаляет zero-width и управляющие символы, шумовые знаки, ложные переносы строк, лишние пробелы, пробелы перед знаками препинания и лишние пробелы внутри скобок. После очистки вызывается should_mark_for_review, который помечает блок для ручной проверки, если текст пустой, confidence ниже порога TEXT_BLOCK_REVIEW_CONFIDENCE, длина текста меньше TEXT_BLOCK_MIN_LENGTH либо фактически использовался резервный backend вместо PaddleOCR.",
    )

    add_table(
        doc,
        "Таблица 6 — Основные правила обработки и контроля качества OCR-блоков",
        ["Этап", "Правило реализации", "Практический эффект"],
        [
            [
                "Сортировка routed-блоков",
                "sort_blocks учитывает reading_order, затем координаты y и x, после чего block_id.",
                "Сохраняется стабильный порядок чтения страницы при повторных прогонах.",
            ],
            [
                "Ограничение и вырезание bbox",
                "clamp_bbox не позволяет выйти за границы изображения; crop_block работает только с корректными координатами.",
                "Исключаются ошибки вырезания и пустые фрагменты из-за некорректной геометрии.",
            ],
            [
                "OCR по двум API",
                "Сначала вызывается predict, затем при необходимости ocr.",
                "Модуль совместим с несколькими поколениями PaddleOCR.",
            ],
            [
                "Fallback по seed_text",
                "При недоступности PaddleOCR используется seed_text и seed_confidence из layout-этапа.",
                "Конвейер продолжает работу и не теряет содержимое text-блоков.",
            ],
            [
                "Очистка текста",
                "postprocess_text устраняет технические символы, ложные переносы и пробельные артефакты.",
                "Снижается число OCR-ошибок, мешающих дальнейшей сборке page_text.",
            ],
            [
                "Маркировка needs_review",
                f"Блок помечается при confidence < {manifest['review_threshold']:.2f}, длине < {manifest['min_length']} или при использовании fallback backend.",
                "Результат становится пригодным для выборочной экспертной верификации.",
            ],
            [
                "Сохранение результата",
                "PageContent сериализуется в /tmp/text-block-results и помещается в ограниченный memory-cache.",
                "Итог можно быстро просматривать через API и веб-интерфейс без повторного OCR.",
            ],
        ],
        [2.9, 6.0, 6.9],
        font_size=9.6,
    )

    add_algorithm(
        doc,
        "Алгоритм 3 — Последовательность OCR-обработки routed-блоков",
        [
            "Вход: page_image, routed_blocks",
            "Выход: page_content",
            "1. image <- LoadPageImage(page_image)",
            "2. ordered_blocks <- SortBlocks(routed_blocks)",
            "3. Для каждого block в ordered_blocks:",
            "   3.1. Если block.route_to = text_pipeline, выполнить crop <- CropBlock(image, block.bbox)",
            "   3.2. Попытаться получить ocr_result <- PaddleOCR(crop) через predict() или ocr()",
            "   3.3. При ошибке OCR использовать ocr_result <- FallbackTextResult(block)",
            "   3.4. processed_text <- PostprocessText(ocr_result.text)",
            "   3.5. review <- ShouldMarkForReview(processed_text, ocr_result.confidence) или backend != paddleocr",
            "   3.6. Сформировать ProcessedBlock с content, confidence, crop_data_url, ocr_backend",
            "   3.7. Иначе сохранить нетекстовый блок без повторного OCR, но с исходным route_to",
            "4. page_content <- BuildPageContent(processed_blocks)",
            "5. SaveResult(page_content) и CachePageContent(page_content)",
            "6. Вернуть page_content",
        ],
    )

    add_paragraph(
        doc,
        "После поэлементной обработки модуль не просто возвращает список распознанных блоков, а собирает полноценное представление страницы. В объекте PageContent одновременно формируются page_text, readable_text и LaTeX-preview. При построении page_text учитываются только text_pipeline-блоки, а формульные области заменяются конфигурируемым placeholder-ом "
        + manifest["formula_placeholder"]
        + ". Благодаря этому следующий этап конвейера получает и компактный линейный текст, и более богатое структурное представление, пригодное для интерфейса и последующих сервисов.",
    )
    add_paragraph(
        doc,
        f"Для демонстрации работы модуля был выполнен реальный прогон по тестовой странице tome1_p02_orig.png. На ней сформировано {page_two['text_block_count']} текстовых routed-блоков, а среднее значение confidence seed-результатов составило {page_two['average_confidence']:.3f}. Поскольку в текущем окружении отсутствует backend paddle, все блоки были обработаны по резервному сценарию surya_seed и помечены для ручной проверки. Одновременно на странице tome1_p01_orig.png первый блок вернул пустое содержимое при confidence {page_one['blocks'][0]['confidence']:.3f}, что демонстрирует корректную работу логики контроля качества для коротких и шумовых областей.",
    )

    add_picture(
        doc,
        ARTIFACTS_DIR / "ocr_block_examples.png",
        width_cm=15.8,
        caption="Рисунок 7 — Примеры crop-областей и итогового содержимого блоков после выполнения process_text_blocks",
    )

    add_table(
        doc,
        "Таблица 7 — Примеры результатов OCR-обработки для тестовых страниц",
        ["Блок", "Краткое содержимое", "Confidence", "Backend", "Причина проверки"],
        [
            [
                row["block_id"],
                row["content"],
                f"{row['confidence']:.3f}",
                row["backend"],
                row["review_reason"],
            ]
            for row in selected_rows
        ],
        [2.8, 6.8, 1.7, 2.3, 2.2],
        font_size=9.0,
    )

    add_paragraph(
        doc,
        "Таким образом, реализованный модуль OCR текстовых блоков решает сразу несколько задач: сохраняет порядок чтения routed-областей, выполняет распознавание или контролируемый fallback, очищает текст от типичных OCR-артефактов, формирует признаки ручной проверки и собирает единый объект PageContent для следующего этапа обработки. Наличие crop-артефактов, поля ocr_backend, JSON-сериализации и кэширования делает модуль не только рабочим компонентом конвейера, но и удобным инструментом для отладки, визуальной верификации и интеграции с пользовательским интерфейсом.",
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


if __name__ == "__main__":
    main()
